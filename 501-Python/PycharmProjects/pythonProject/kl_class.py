import os  # 用于和系统相关的操作

import configparser  # 用于配置文件ini的操作

import pandas as pd  # 用于pandas操作
from pandas.api.types import CategoricalDtype  # 用于DataFrame排序用。
import matplotlib.pyplot as plt  # 用于pandas 绘图
from pylab import mpl  # 解决中文问题

import pymysql  # 用于操作MySQL数据库
from sqlalchemy import create_engine  # 用于解决pandas数据库连接问题


import docx  # 用于操作word文档
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

import re  # 用于正则表达式

import time  # 用于时间操作
from dateutil.relativedelta import relativedelta  # 用于获取时间间隔
import datetime  # 用于日期操作
import calendar  # 用于日历操作
from dateutil import rrule  # 用于时间间隔


# 用于获取电脑名称
def get_pc_name():
    return os.environ['COMPUTERNAME']


# 用于获取配置文件路径，区分办公室和个人电脑
def ini_path():
    if get_pc_name() == 'LAPTOP-HF9P6H1P':
        in_pt = r'D:\JGY\600-Data\001-ini配置文件\孔令配置.ini'
    else:
        in_pt = r'Z:\600数据库\000配置文件\昆明局安监部配置文件.ini'
    return in_pt


# 实现删除list中指定list
def list_del_list(o_list, d_list):
    r_list = []
    for i in o_list:
        if i not in d_list:
            r_list.append(i)
    return r_list


# 按照指定列及list排序，可以实现Series和DataFrame,也可以实现排序list少于排序列值的排序。
def sort_list(obj, list_sort, column=None):
    if isinstance(obj, pd.Series):
        df=pd.DataFrame(obj)
        df = df.reset_index()
        df.columns = ['index', 'values']
        df1 = df[df['index'].isin(list_sort)].copy()
        df2 = df[~df['index'].isin(list_sort)].copy()
        df1['index'] = df1['index'].astype('category')
        df1['index'].cat.set_categories(list_sort, inplace=True)
        df1.sort_values('index', ascending=True, inplace=True)
        df = pd.concat([df1, df2])
        object_ = pd.Series(df['values'].values, index=df['index'])
    elif isinstance(obj, pd.DataFrame):
        for i in range(len(column)):
            list_sort[i] = list_sort[i] + list(set(obj[column[i]]).difference(set(list_sort[i])))
        for i in range(len(column)):
            cat_order = CategoricalDtype(
                list_sort[i],
                ordered=True
            )
            obj[column[i]] = obj[column[i]].astype(cat_order)
        object_ = obj.sort_values(column, axis=0, ascending=[True] * (len(column)))
    return object_


# 实现Excel文件更新到MySQL中
class xlsxMysql(object):
    # __doc__内容
    """
    孔令的办公自动化库，将excel更新到mysql数据库。
    """
    def __init__(self, path, sheet_name):
        self.path = path
        self.sheet_name = sheet_name

    def read_xlsx(self):
        if self.path.endswith('.xls'):
            df_xl = pd.read_excel(self.path, sheet_name=self.sheet_name)
        elif self.path.endswith('.xlsx'):
            df_xl = pd.read_excel(self.path, sheet_name=self.sheet_name, engine='openpyxl')
        else:
            print('你选择的文件不是excel文件，请选后缀为.xls或.xlsx文件')
        return df_xl
    def to_mysql(self,table_name):
        dfsql = dfMysql()
        dfsql.replace_mysql(self.read_xlsx(), table_name)


# 实现DataFrame 更新到 MySQL
class dfMysql():
    def __init__(self):
        # 读取配置文件ini
        config = configparser.ConfigParser(interpolation=configparser.ExtendedInterpolation())
        config.read(ini_path(), encoding='utf-8')
        self.connect = pymysql.connect(
            host=config.get('MySQL', 'host'),
            port=int(config.get('MySQL', 'port')),
            user=config.get('MySQL', 'user'),
            passwd=config.get('MySQL', 'passwd'),
            db=config.get('MySQL', 'db'),
            charset=config.get('MySQL', 'charset')
        )
        con = "mysql+pymysql://{0}:{1}@{2}:{3}/{4}?chart{5}".format(config.get('MySQL', 'user'),
                                                                    config.get('MySQL', 'passwd'),
                                                                    config.get('MySQL', 'host'),
                                                                    config.get('MySQL', 'port'),
                                                                    config.get('MySQL', 'db'),
                                                                    config.get('MySQL', 'charset')
                                                                    )
        self.engine = create_engine(con)
    # 查询mysql数据
    def select_mysql(self, sql):
        df = pd.read_sql(sql, self.engine)
        return df
    # 删除一条数据
    def delete_mysql(self,sql):
        cn = self.connect.cursor()
        cn.execute(sql)
        self.connect.commit()
        cn.close()
        self.connect.close()
    # 追加或更新
    def replace_mysql(self,df,table_name):
        try:
            # 执行SQL语句
            df.to_sql('temp', self.engine, if_exists='replace', index=False)  # 把新数据写入 temp 临时表
            # 替换数据的语句
            cn = self.connect.cursor()
            args1 = f" REPLACE INTO {table_name} SELECT * FROM temp "
            cn.execute(args1)
            args2 = " DROP Table If Exists temp"  # 把临时表删除
            cn.execute(args2)
            # 提交到数据库执行
            self.connect.commit()
            cn.close()
            self.connect.close()
            print('更新成功')

        except Exception as e:
            print('更新失败')
            print(e)
            # 发生错误时回滚
            self.connect.rollback()
            # 关闭数据库连接
            cn.close()
            self.connect.close()


class parsingDf():
    def __init__(self):
        pass

    def series_tex(self,sr, unit):
        #print(sr)
        text = ''
        for i in range(len(sr)):
            if i != len(sr) - 1:
                text += sr.index[i] +' '+ str(sr[i]) + unit + '，'
            else:
                text += sr.index[i]+' ' + str(sr[i]) + unit + '。'

        return text

    def df_iter(self,df):
        #print(df)
        y = []
        j=1
        for row in df.iterrows():
            x = []
            for i in range(len(row[1])):
                x.append(row[1].index[i] + digital_to_chinese(j) + '：' + str(row[1][row[1].index[i]]) + '\n')
            y.append(''.join(x))
            j+=1
        #print(y)
        return y

    def null_item(self,df):
        null_ = df.isnull().sum()
        null_.sort_values(ascending=False, inplace=True)
        null_text = self.series_tex(null_, '项')
        return null_,null_text
    def df_item(self,df,column):
        item=df[column].value_counts()
        item_text = self.series_tex(item, '项')
        return item,item_text
    def df_fig(self,df,fig_class):
        mpl.rcParams['font.sans-serif'] = ['Microsoft YaHei']
        # 解决图表负号显示不正确问题
        plt.rcParams['axes.unicode_minus'] = False
        ax = df.plot(kind=fig_class,
                     legend=False,
                     figsize=(8,5),
                     rot=20
                     #title='Pie of Weather in London'
                     )
        fig = ax.get_figure()
        if get_pc_name() == 'LAPTOP-HF9P6H1P':
            # 个人
            fig.savefig(r'D:\JGY\600-Data\006-temporary临时文件\fig.png')
        else:
            # 单位
            fig.savefig(r'D:\out\fig.png')
        plt.clf()
        return


class oas_docx():
    # __doc__内容

    '''

    孔令的办公自动化库，这是word自动化的类，主要能建立docx文档，并且可以对文档
    进行设置，简化python-docx的重复设置，只需应用就可以。


    '''

    # 构造函数
    def __init__(self,
                 d_path=None,
                 d_styles_name='Normal',
                 f_name='宋体',
                 p_height=29.7, p_width=21,
                 l_margin=3.17, r_margin=3.17, t_margin=2.54, b_margin=2.54):

        self.d_path = d_path
        self.d_styles_name = d_styles_name
        self.f_name = f_name
        self.p_height = p_height
        self.p_width = p_width
        self.l_margin = l_margin
        self.r_margin = r_margin
        self.t_margin = t_margin
        self.b_margin = b_margin
        self.table = None
        # 写入docx
        self.doc = docx.Document(d_path)
        self.doc.styles[d_styles_name].font.name = f_name
        self.doc.styles[d_styles_name]._element.rPr.rFonts.set(qn('w:eastAsia'), f_name)

        sec = self.doc.sections
        # 文档页边距设置
        # 获取、设置页面边距
        sec0 = sec[0]  # 获取章节对象
        sec0.page_height = Cm(p_height)
        sec0.page_width = Cm(p_width)
        # 设置页面的边距：
        sec0.left_margin = Cm(l_margin)
        sec0.right_margin = Cm(r_margin)
        sec0.top_margin = Cm(t_margin)
        sec0.bottom_margin = Cm(b_margin)

    def hd(self, text='', level=1, font_name='黑体', alig='center', f_indent=32, font_size=16, l_space=28):
        self.head = self.doc.add_heading('', level)
        run = self.head.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        if alig == 'left':
            self.head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            self.head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            self.head.alignment.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def par(self, text: str, bold=False,font_name='宋体', alig='left', f_indent=32, font_size=16, l_space=28):
        p = self.doc.add_paragraph()  # 插入段落
        p_format = p.paragraph_format
        # 设置首行缩进量
        p_format.first_line_indent = Pt(f_indent)
        # 设置行间距,若为整数则是数值，若为小数则为N倍行距
        if isinstance(l_space, int):
            p_format.line_spacing = Pt(l_space)
        elif isinstance(l_space, float):
            p_format.line_spacing = l_space

        # 段前间距
        p_format.space_before = Pt(0)
        # 段后间距
        p_format.space_after = Pt(0)

        # 设置段落水平对齐方式
        if alig == 'left':
            p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)

        #run.font.bold = True
        font = run.font
        font.bold = bold
        # 设置中文字体，需要两语句同时使用。
        font.name = font_name


        font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 设置字体大小
        font.size = Pt(font_size)
        return p

    def pic(self, path: str, alig='left', wth=1.25):
        # 插入图片
        p = self.doc.add_paragraph()  # 插入段落
        paragraph_format = p.paragraph_format
        # 设置段落水平对齐方式
        if alig == 'left':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(path, width=Inches(wth))
        return p

    def tb(self, row, column, style='Table Grid'):
        self.table = self.doc.add_table(row, column, style='Table Grid')

        # self.table.cell(0,3).width=Cm(5)
        # self.table.rows[0].height=Cm(1.5)
        # self.table.rows[1].height=Cm(1.5)
        # self.table.rows[2].height=Cm(5)
        # self.table.rows[3].height=Cm(5)
        # self.table.rows[4].height=Cm(5)
        # table.cell(3,0).height=Cm(5)

    def tb_cell(self, text: str, row: int, column: int, alig='left', v_alig='center', f_name='宋体', f_size=10.5):
        tb_cell = self.table.cell(row, column)
        # 设置单元格垂直对齐方式
        if alig == 'left':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        elif alig == 'center':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif alig == 'right':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        # 设置单元格水平对齐方式
        cell_par = self.table.cell(row, column).paragraphs[0]  # 获取到对象
        if alig == 'left':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 增加run和文字
        font = cell_par.add_run(text).font
        # 设置字体
        font.name = f_name
        font.element.rPr.rFonts.set(qn('w:eastAsia'), f_name)
        font.size = Pt(f_size)
        return tb_cell

    def save_docx(self, docx_name: str):
        self.doc.save(docx_name)  # 保存文档


class nDay():
    '''
    此类用于获取给定日期所在月第一天和最后一天，或任意日的间隔一个月的开始和结束日期；
    实例变量只需要给定任意日期。
    the_n_day方法参数n为需要返回的开始日,m为需要返回的间隔月。
    例如：返回给定日期datetime所在月的第一天和最后一天
    d=nDay（datetime）
    d.the_n
    '''

    def __init__(self, date_time: datetime.datetime):
        self.date_time = date_time

    # 获取输入日期所在月份的第一天，最后一天和第n天，依据m可以是输入日期的前后n个月
    def get_n_day(self, n=1, m=0):
        # this_month_start = datetime.datetime(self.date_time.year, self.date_time.month, 1)
        this_month_nday = datetime.datetime(self.date_time.year, self.date_time.month, n)  # +datetime.timedelta(days=n)
        this_month_end = datetime.datetime(self.date_time.year, self.date_time.month,
                                           calendar.monthrange(self.date_time.year, self.date_time.month)[1])
        # n_month_start=this_month_start +relativedelta(months=m)
        n_month_end = this_month_end + relativedelta(months=m)
        n_month_nday = this_month_nday + relativedelta(months=m)

        return n_month_nday, n_month_end

    # 获取输入日期所在周的第一天，最后一天和第n天，依据m可以是输入日期的前后n个周
    def get_current_week(self, date_time, n):
        monday, sunday = date_time, date_time
        one_day = datetime.timedelta(days=1)
        while monday.weekday() != 0:
            monday -= one_day
        while sunday.weekday() != 6:
            sunday += one_day
        # 返回当前的星期一和星期天的日期
        week_n = monday + datetime.timedelta(days=n)

        return monday, sunday, week_n


class periodConvert():
    def __init__(self, dataframe, start, end, interval):
        self.df = dataframe
        self.start = start
        self.end = end
        self.interval = interval

    def p_split(self):
        df_scope = len(self.df)
        for i in range(df_scope):
            if isinstance(self.df[self.start][i], datetime.datetime) and isinstance(self.df[self.end][i],
                                                                                    datetime.datetime):
                d_start = self.df[self.start][i]
                d_end = self.df[self.end][i]
                if self.interval == 'm':
                    delta = rrule.rrule(rrule.MONTHLY, dtstart=d_start, until=d_end).count()
                    loop_delta = 0
                    if delta > 1:
                        loop_delta = delta
                        for j in range(loop_delta):
                            df_scope = j + df_scope
                            self.df.loc[df_scope] = self.df.loc[i]
                            this_month_start = datetime.datetime(self.df[self.start][df_scope].year,
                                                                 self.df[self.start][df_scope].month, 1)
                            this_month_end = datetime.datetime(self.df[self.start][df_scope].year,
                                                               self.df[self.start][df_scope].month,
                                                               calendar.monthrange(self.df[self.start][df_scope].year,
                                                                                   self.df[self.start][df_scope].month)[
                                                                   1])
                            if j == 0:
                                self.df.loc[df_scope, self.end] = this_month_end
                            elif j == loop_delta:
                                self.df.loc[df_scope, self.start] = this_month_start

                            else:
                                self.df.loc[df_scope, self.start] = this_month_start + relativedelta(months=+j)
                                self.df.loc[df_scope, self.end] = this_month_end + relativedelta(months=+j)

                    df_scope = df_scope + 1
                    self.df.drop(index=[i], inplace=True)

    def p_merge(self, append_column: list):
        l_temp = list(self.df)
        for i in append_column:
            l_temp.remove(i)
        temp_df = self.df.duplicated(l_temp)


def digital_to_chinese(num):
    num_dict = {'1': '一', '2': '二', '3': '三', '4': '四', '5': '五', '6': '六', '7': '七', '8': '八', '9': '九', '0': '〇', }
    index_dict = {1: '', 2: '十', 3: '百', 4: '千', 5: '万', 6: '十', 7: '百', 8: '千', 9: '亿'}
    nums = list(str(num))
    nums_index = [x for x in range(1, len(nums) + 1)][-1::-1]
    str_ = ''
    for index, item in enumerate(nums):
        str_ = "".join((str_, num_dict[item], index_dict[nums_index[index]]))

    str_ = re.sub("〇[十百千〇]*", "〇", str_)
    str_ = re.sub("〇万", "万", str_)
    str_ = re.sub("亿万", "亿〇", str_)
    str_ = re.sub("〇〇", "〇", str_)
    str_ = re.sub("〇\\b", "", str_)
    return str_


# 获取输入日期所在月份的第一天，最后一天和第n天，依据m可以是输入日期的前后n个月
def the_n_day(date_time, n: int, m: int):
    this_month_start = datetime.datetime(date_time.year, date_time.month, 1)
    this_month_nday = datetime.datetime(date_time.year, date_time.month, 1) + datetime.timedelta(days=n)
    this_month_end = datetime.datetime(date_time.year, date_time.month,
                                       calendar.monthrange(date_time.year, date_time.month)[1])
    n_month_start = this_month_start + relativedelta(months=m)
    n_month_end = this_month_end + relativedelta(months=m)
    n_month_nday = this_month_nday + relativedelta(months=m)

    return n_month_start, n_month_end, n_month_nday


# 获取输入日期所在周的第一天，最后一天和第n天，依据m可以是输入日期的前后n个周
def get_current_week(date_time, n):
    monday, sunday = date_time, date_time
    one_day = datetime.timedelta(days=1)
    while monday.weekday() != 0:
        monday -= one_day
    while sunday.weekday() != 6:
        sunday += one_day
    # 返回当前的星期一和星期天的日期
    week_n = monday + datetime.timedelta(days=n)

    return monday, sunday, week_n


# 将用sign符号分隔的字符串string，进行逆序操作
# 例如：string 1/2/3，sign /，reverse_item后 输出 3/2/1
def reverse_item(string: str, sign: str):
    # start=time.time()
    # 将字符串按照sign转成list
    list_str = str.split(string, sign)
    # 将列表中的值进行逆序操作，reverse直接改变list值
    list_str.reverse()
    # 从新用sign符号将列表进行合并
    string_rv = sign.join(list_str)
    # print('reverse_item耗时%d',time.time()-start)
    return string_rv


# 将dataframe中sp_cloc列中用sign分隔的内容分列成n条
def split_record(df: pd.DataFrame, sp_cloc: str, sign: str):
    # 记录原始dataframe的长度，并作为后面追加数据的基准
    df.reset_index(inplace=True,drop=True)
    y = len(df)
    # 记录原始dataframe的长度
    len_df = len(df)
    for i in range(len(df)):
        # 将需要分列的列数据按照sign分列出来
        #print(i,df.loc[i, sp_cloc])
        list_str = str.split(df.loc[i, sp_cloc], sign)
        # 依据分列出来数据增加行
        for j in range(len(list_str)):
            y = y + j
            # 复制i行数据
            df.loc[y] = df.loc[i]
            # 将y行数据sp_cloc列数据设置为list_str中数据
            df.loc[y, sp_cloc] = list_str[j]
        y = y + 1
    # 取出新增加数据
    df = df.tail(len(df) - len_df)
    # 重置index
    df = df.reset_index(drop=True)
    return df


# 将某列以外都相同的行以','进行合并
def plus_record(dfp: pd.DataFrame, pr_cloc: str, sign: str):
    # 获取dfp的列名
    list_cl = list(dfp)
    # 去除需要合并列的列名
    list_cl.remove(pr_cloc)
    # 按照去除合并列名进行排序
    dfp.sort_values(list_cl, inplace=True)
    # 按照排序后行，重新设置index
    dfp = dfp.reset_index(drop=True)
    # 复制dfp，用以处理
    dfpc = dfp.copy()
    # 去除合并列
    dfpc.drop([pr_cloc], axis=1, inplace=True)
    # 进行查重处理
    list_dp = dfpc.duplicated()
    # 查找重复分界点
    x = list_dp[list_dp.isin([False])].index
    # 因为没有找到index插入的方法，将分界点index转为list
    list_x = []
    for q in range(len(x)):
        list_x.append(x[q])
    # 主要用于加入最后一条记录index
    list_x.append(len(dfp))

    # print(list_x)
    # x.append(int64(len(dfp)))
    yn = []
    # 循环获取重复记录段数据
    for i in range(len(list_x) - 1):
        # 判断是否有需要合并项
        if (list_x[i + 1] - list_x[i]) > 1:
            # 若有序号间隔大于1，则进入循环
            for j in range(list_x[i + 1] - list_x[i]):
                # 取出需要合并数据，形成list
                yn.append(dfp.loc[list_x[i] + j, pr_cloc])
            # 将list合并成以sign为分隔字符串。
            y = sign.join(yn)
            # 将字符串赋给dfp第一列
            dfp.loc[list_x[i], pr_cloc] = y
            # 删除多余项目
            for k in range(list_x[i + 1] - list_x[i] - 1):
                dfp.drop(list_x[i] + 1 + k, axis=0, inplace=True)
            # 清空记录list
            yn = []
    # 重置index
    dfp = dfp.reset_index(drop=True)

    return dfp


def append_xlsx(dirs: str, sheet_nm: str, header_row: int, skipfoot_row: int):
    # 读取文件夹下所有文件名
    myList = os.listdir(dirs)
    # 打印文件个数
    print('合并Excel文件个数：', len(myList))
    # 读取第一个文件
    inExcelFile = dirs + '\\' + myList[0]
    df1 = pd.read_excel(inExcelFile, sheet_name=sheet_nm, header=header_row, skipfooter=skipfoot_row)

    print('读取第一个Excel文件成功。')
    # 循环追加excel文件
    for i in range(len(myList) - 1):
        inExcelFile = dirs + '\\' + myList[i + 1]
        df2 = pd.read_excel(inExcelFile, sheet_name=sheet_nm, header=header_row, skipfooter=skipfoot_row)
        # df2=read_excel(inExcelFile,sheet_name=exSh,header=int(exSl)-1)
        print('读取第', i + 1, '个Excel文件成功。')
        df1 = df1.append(df2, ignore_index=True)
        print('合并第', i + 1, '个和第', i + 2, '个文件成功。')
    print('合并', i + 2, '个文件成功。')
    # 去除重复项目
    df1 = df1.drop_duplicates()
    print('去除重复项成功。')
    # 去除空白行
    df1 = df1.dropna(axis=0, how='all', inplace=False)
    print('去除空白行成功。')
    return df1


# 将doc文档转为docx文档
from win32com import client


def doc_to_docx(path):
    if os.path.splitext(path)[1] == ".doc":
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(path)  # 目标路径下的文件
        doc.SaveAs(os.path.splitext(path)[0] + ".docx", 16)  # 转化后路径下的文件
        doc.Close()
        word.Quit()
        os.remove(path)


# 创建读取doc文档内容的函数，去除格式
import docx


def readDocx(docName):
    fullText = []
    doc = docx.Document(docName)
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)


# 如果文件夹不存在创建文件夹
def Makedir(path):
    folder = os.path.exists(path)
    if (not folder):
        os.makedirs(path)