# -*- coding: utf-8 -*-
"""
Created on Mon Oct 19 10:34:57 2020

@author: KongLing

"""


import docx
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

import pandas as pd


class OasDocx(object):
    # __doc__内容
    """
    孔令的办公自动化库，这是word自动化的类，主要能建立docx文档，并且可以对文档
    进行设置，简化python-docx的重复设置，只需应用就可以。
    更新时间2021-2-19，版本v0.69版本
    """

    # 构造函数
    def __init__(self,
                 d_path=None,
                 d_styles_name='Normal',
                 f_name='宋体',
                 p_height=29.7, p_width=21,
                 l_margin=3.17, r_margin=3.17, t_margin=2.54, b_margin=2.54):

        self.d_path = d_path
        self.d_styles_name = d_styles_name
        self.f_name = f_name
        self.p_height = p_height
        self.p_width = p_width
        self.l_margin = l_margin
        self.r_margin = r_margin
        self.t_margin = t_margin
        self.b_margin = b_margin
        self.head = None
        self.table = None
        # 写入docx
        self.doc = docx.Document(d_path)
        self.doc.styles[d_styles_name].font.name = f_name
        self.doc.styles[d_styles_name]._element.rPr.rFonts.set(qn('w:eastAsia'), f_name)

        sec = self.doc.sections
        # 文档页边距设置
        # 获取、设置页面边距
        sec0 = sec[0]  # 获取章节对象
        sec0.page_height = Cm(p_height)
        sec0.page_width = Cm(p_width)
        # 设置页面的边距：
        sec0.left_margin = Cm(l_margin)
        sec0.right_margin = Cm(r_margin)
        sec0.top_margin = Cm(t_margin)
        sec0.bottom_margin = Cm(b_margin)

    # 插入标题
    def hd(self, text='', level=1, font_name='黑体', a_lig='center', f_indent=32, font_size=16, l_space=28):
        self.head = self.doc.add_heading('', level)
        run = self.head.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        if a_lig == 'left':
            self.head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif a_lig == 'center':
            self.head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif a_lig == 'right':
            self.head.alignment.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 插入段落
    def par(self, text: str, font_name='宋体', bold=False, a_lig='left', f_indent=32, font_size=16, l_space=28):
        p = self.doc.add_paragraph()  # 插入段落
        p_format = p.paragraph_format
        # 设置首行缩进量
        p_format.first_line_indent = Pt(f_indent)
        # 设置行间距,若为整数则是数值，若为小数则为N倍行距
        if isinstance(l_space, int):
            p_format.line_spacing = Pt(l_space)
        elif isinstance(l_space, float):
            p_format.line_spacing = l_space

        # 段前间距
        p_format.space_before = Pt(0)
        # 段后间距
        p_format.space_after = Pt(0)

        # 设置段落水平对齐方式
        if a_lig == 'left':
            p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif a_lig == 'center':
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif a_lig == 'right':
            p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.bold = bold
        font = run.font
        # 设置中文字体，需要两语句同时使用。
        font.name = font_name
        font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 设置字体大小
        font.size = Pt(font_size)
        return p

    # 插入图表
    def pic(self, path: str, a_lig='left', wth=1.25):
        # 插入图片
        p = self.doc.add_paragraph()  # 插入段落
        paragraph_format = p.paragraph_format
        # 设置段落水平对齐方式
        if a_lig == 'left':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif a_lig == 'center':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif a_lig == 'right':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(path, width=Inches(wth))
        return p

    # 插入表格，目前还没有完善可以自动优化调整表格间距功能
    def tb(self, row, column, style='Table Grid'):
        self.table = self.doc.add_table(row, column, style=style)
        # self.table.cell(0,3).width=Cm(5)
        # self.table.rows[0].height=Cm(1.5)
        # self.table.rows[1].height=Cm(1.5)
        # self.table.rows[2].height=Cm(5)
        # self.table.rows[3].height=Cm(5)
        # self.table.rows[4].height=Cm(5)
        # table.cell(3,0).height=Cm(5)

    # 插入表格单元格，主要为了可以调整表格中文字格式
    def tb_cell(self, text: str, row: int, column: int, a_lig='left', v_a_lig='center', f_name='宋体', f_size=10.5):
        tb_cell = self.table.cell(row, column)
        # 设置单元格垂直对齐方式
        if a_lig == 'left':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        elif a_lig == 'center':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif a_lig == 'right':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        # 设置单元格水平对齐方式
        cell_par = self.table.cell(row, column).paragraphs[0]  # 获取到对象
        if a_lig == 'left':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif a_lig == 'center':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif a_lig == 'right':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 增加run和文字
        font = cell_par.add_run(text).font
        # 设置字体
        font.name = f_name
        font.element.rPr.rFonts.set(qn('w:eastAsia'), f_name)
        font.size = Pt(f_size)
        return tb_cell

    # 插入表格，可以修改表格格式
    def d_table(self, d_table_df):
        # 确保df后面引用index不出错
        d_table_df.reset_index(drop=True, inplace=True)
        rows = d_table_df.shape[0] + 1
        cols = d_table_df.shape[1]
        self.tb(rows, cols)
        head = list(d_table_df)
        for i in range(rows):
            for j in range(cols):
                if i == 0:
                    self.tb_cell(str(head[j]), i, j, f_name='黑体')  # 可以更改表头字体等
                else:
                    self.tb_cell(str(d_table_df.iloc[i - 1, j]), i, j)  # 可以更改表格内容和字体

    def save_docx(self, docx_name: str):
        self.doc.save(docx_name)  # 保存文档


if __name__ == '__main__':
    df = pd.DataFrame(
        [
            ['A', 'B', 'C', pd.to_datetime('2020-12-31'), pd.to_datetime('2021-1-8')],
            ['D', 'E', 'F', pd.to_datetime('2021-2-26'), pd.to_datetime('2021-3-7')],
            ['C', 'E', 'F', pd.to_datetime('2021-1-20'), pd.to_datetime('2022-3-8')],
        ], columns=['一', '二', '三', '开始日期', '结束日期'])
    od = OasDocx()
    # 文档标题
    od.hd('表格')
    od.par('表格示例')
    od.d_table(df)

    od.save_docx(r'D:\JGY\600-Data\006-temporary临时文件\test.docx')
    print(OasDocx.__doc__)


