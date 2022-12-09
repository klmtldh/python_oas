import os

import configparser
import pandas as pd
import matplotlib.pyplot as plt
from pylab import mpl

import pymysql
from sqlalchemy import create_engine


import docx
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

import re

def get_pc_name():
    return os.environ['COMPUTERNAME']

def ini_path():
    if get_pc_name() == 'LAPTOP-HF9P6H1P':
        ini_path = r'D:\JGY\600-Data\001-ini配置文件\孔令配置.ini'
    else:
        ini_path = r'Z:\600数据库\000配置文件\昆明局安监部配置文件.ini'
    return ini_path

def sort_list(series,list_sort):
    df=pd.DataFrame(series)
    df = df.reset_index()
    df.columns = ['index', 'values']
    df['index']=df['index'].astype('category')
    df['index'].cat.set_categories(list_sort, inplace=True)
    df.sort_values('index', ascending=True, inplace=True)
    #df=df.set_index('index',drop=True)
    s = pd.Series(df['values'].values,index = df['index'])
    return s



class xlsxMysql():
    def __init__(self,path,sheet_name):
        self.path=path
        self.sheet_name=sheet_name

    def read_xlsx(self):
        if self.path.endswith('.xls'):
            df_xl = pd.read_excel(self.path, sheet_name=self.sheet_name)
        elif self.path.endswith('.xlsx'):
            df_xl = pd.read_excel(self.path, sheet_name=self.sheet_name, engine='openpyxl')
        else:
            print('你选择的文件不是excel文件，请选后缀为.xls或.xlsx文件')
        return df_xl
    def to_mysql(self,table_name):
        dfsql = dfMysql()
        dfsql.replace_mysql(self.read_xlsx(), table_name)

class dfMysql():
    def __init__(self):
        # 读取配置文件ini
        config = configparser.ConfigParser(interpolation=configparser.ExtendedInterpolation())
        config.read(ini_path(), encoding='utf-8')
        #print(ini_path())
        self.connect = pymysql.connect(
            host=config.get('MySQL', 'host'),
            port=int(config.get('MySQL', 'port')),
            user=config.get('MySQL', 'user'),
            passwd=config.get('MySQL', 'passwd'),
            db=config.get('MySQL', 'db'),
            charset=config.get('MySQL', 'charset')
        )
        con = 'mysql+pymysql://' + config.get('MySQL', 'user') + ':' \
                      + config.get('MySQL', 'passwd') + '@'+config.get('MySQL', 'host') +':'+ config.get('MySQL', 'port') \
                      + '/' + config.get('MySQL', 'db')+'?'+'charset'+config.get('MySQL', 'charset')

        self.engine = create_engine(con)
    #查询mysql数据
    def select_mysql(self, sql):
        df = pd.read_sql(sql, self.engine)
        return df
    def delete_mysql(self,sql):
        cn = self.connect.cursor()
        cn.execute(sql)
        self.connect.commit()
        cn.close()
        self.connect.close()
    #追加或更新
    def replace_mysql(self,df,table_name):
        try:
            # 执行SQL语句
            #print(df)
            df.to_sql('temp', self.engine, if_exists='replace', index=False)  # 把新数据写入 temp 临时表
            # 替换数据的语句
            cn = self.connect.cursor()
            args1 = f" REPLACE INTO {table_name} SELECT * FROM temp "
            cn.execute(args1)
            args2 = " DROP Table If Exists temp"  # 把临时表删除
            cn.execute(args2)
            # 提交到数据库执行
            self.connect.commit()
            cn.close()
            self.connect.close()
            print('更新成功')

        except Exception as e:
            print('更新失败')
            print(e)
            # 发生错误时回滚
            self.connect.rollback()
            # 关闭数据库连接
            cn.close()
            self.connect.close()

class parsingDf():
    def __init__(self):
        pass

    def series_tex(self,sr, unit):
        #print(sr)
        text = ''
        for i in range(len(sr)):
            if i != len(sr) - 1:
                text += sr.index[i] +' '+ str(sr[i]) + unit + '，'
            else:
                text += sr.index[i]+' ' + str(sr[i]) + unit + '。'

        return text

    def df_iter(self,df):
        #print(df)
        y = []
        j=1
        for row in df.iterrows():
            x = []
            for i in range(len(row[1])):
                x.append(row[1].index[i] + digital_to_chinese(j) + '：' + str(row[1][row[1].index[i]]) + '\n')
            y.append(''.join(x))
            j+=1
        #print(y)
        return y

    def null_item(self,df):
        null_ = df.isnull().sum()
        null_.sort_values(ascending=False, inplace=True)
        null_text = self.series_tex(null_, '项')
        return null_,null_text
    def df_item(self,df,column):
        item=df[column].value_counts()
        item_text = self.series_tex(item, '项')
        return item,item_text
    def df_fig(self,df,fig_class):
        mpl.rcParams['font.sans-serif'] = ['Microsoft YaHei']
        # 解决图表负号显示不正确问题
        plt.rcParams['axes.unicode_minus'] = False
        ax = df.plot(kind=fig_class,
                     legend=False,
                     figsize=(8,5),
                     rot=20
                     #title='Pie of Weather in London'
                     )
        fig = ax.get_figure()
        if get_pc_name() == 'LAPTOP-HF9P6H1P':
            # 个人
            fig.savefig(r'D:\JGY\600-Data\003-out输出文件\工作\fig.png')
        else:
            # 单位
            fig.savefig(r'D:\out\fig.png')
        plt.clf()
        return

class oas_docx():
    # __doc__内容
    '''
    孔令的办公自动化库，这是word自动化的类，主要能建立docx文档，并且可以对文档
    进行设置，简化python-docx的重复设置，只需应用就可以。


    '''

    # 构造函数
    def __init__(self,
                 d_path=None,
                 d_styles_name='Normal',
                 f_name='宋体',
                 p_height=29.7, p_width=21,
                 l_margin=3.17, r_margin=3.17, t_margin=2.54, b_margin=2.54):

        self.d_path = d_path
        self.d_styles_name = d_styles_name
        self.f_name = f_name
        self.p_height = p_height
        self.p_width = p_width
        self.l_margin = l_margin
        self.r_margin = r_margin
        self.t_margin = t_margin
        self.b_margin = b_margin
        self.table = None
        # 写入docx
        self.doc = docx.Document(d_path)
        self.doc.styles[d_styles_name].font.name = f_name
        self.doc.styles[d_styles_name]._element.rPr.rFonts.set(qn('w:eastAsia'), f_name)

        sec = self.doc.sections
        # 文档页边距设置
        # 获取、设置页面边距
        sec0 = sec[0]  # 获取章节对象
        sec0.page_height = Cm(p_height)
        sec0.page_width = Cm(p_width)
        # 设置页面的边距：
        sec0.left_margin = Cm(l_margin)
        sec0.right_margin = Cm(r_margin)
        sec0.top_margin = Cm(t_margin)
        sec0.bottom_margin = Cm(b_margin)

    def hd(self, text='', level=1, font_name='黑体', alig='center', f_indent=32, font_size=16, l_space=28):
        self.head = self.doc.add_heading('', level)
        run = self.head.add_run(text)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = True
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        if alig == 'left':
            self.head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            self.head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            self.head.alignment.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def par(self, text: str, bold=False,font_name='宋体', alig='left', f_indent=32, font_size=16, l_space=28):
        p = self.doc.add_paragraph()  # 插入段落
        p_format = p.paragraph_format
        # 设置首行缩进量
        p_format.first_line_indent = Pt(f_indent)
        # 设置行间距,若为整数则是数值，若为小数则为N倍行距
        if isinstance(l_space, int):
            p_format.line_spacing = Pt(l_space)
        elif isinstance(l_space, float):
            p_format.line_spacing = l_space

        # 段前间距
        p_format.space_before = Pt(0)
        # 段后间距
        p_format.space_after = Pt(0)

        # 设置段落水平对齐方式
        if alig == 'left':
            p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)

        #run.font.bold = True
        font = run.font
        font.bold = bold
        # 设置中文字体，需要两语句同时使用。
        font.name = font_name


        font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 设置字体大小
        font.size = Pt(font_size)
        return p

    def pic(self, path: str, alig='left', wth=1.25):
        # 插入图片
        p = self.doc.add_paragraph()  # 插入段落
        paragraph_format = p.paragraph_format
        # 设置段落水平对齐方式
        if alig == 'left':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(path, width=Inches(wth))
        return p

    def tb(self, row, column, style='Table Grid'):
        self.table = self.doc.add_table(row, column, style='Table Grid')

        # self.table.cell(0,3).width=Cm(5)
        # self.table.rows[0].height=Cm(1.5)
        # self.table.rows[1].height=Cm(1.5)
        # self.table.rows[2].height=Cm(5)
        # self.table.rows[3].height=Cm(5)
        # self.table.rows[4].height=Cm(5)
        # table.cell(3,0).height=Cm(5)

    def tb_cell(self, text: str, row: int, column: int, alig='left', v_alig='center', f_name='宋体', f_size=10.5):
        tb_cell = self.table.cell(row, column)
        # 设置单元格垂直对齐方式
        if alig == 'left':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        elif alig == 'center':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif alig == 'right':
            tb_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        # 设置单元格水平对齐方式
        cell_par = self.table.cell(row, column).paragraphs[0]  # 获取到对象
        if alig == 'left':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alig == 'center':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alig == 'right':
            cell_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 增加run和文字
        font = cell_par.add_run(text).font
        # 设置字体
        font.name = f_name
        font.element.rPr.rFonts.set(qn('w:eastAsia'), f_name)
        font.size = Pt(f_size)
        return tb_cell

    def save_docx(self, docx_name: str):
        self.doc.save(docx_name)  # 保存文档





def digital_to_chinese(num):
    num_dict = {'1': '一', '2': '二', '3': '三', '4': '四', '5': '五', '6': '六', '7': '七', '8': '八', '9': '九', '0': '〇', }
    index_dict = {1: '', 2: '十', 3: '百', 4: '千', 5: '万', 6: '十', 7: '百', 8: '千', 9: '亿'}
    nums = list(str(num))
    nums_index = [x for x in range(1, len(nums) + 1)][-1::-1]
    str_ = ''
    for index, item in enumerate(nums):
        str_ = "".join((str_, num_dict[item], index_dict[nums_index[index]]))

    str_ = re.sub("〇[十百千〇]*", "〇", str_)
    str_ = re.sub("〇万", "万", str_)
    str_ = re.sub("亿万", "亿〇", str_)
    str_ = re.sub("〇〇", "〇", str_)
    str_ = re.sub("〇\\b", "", str_)
    return str_